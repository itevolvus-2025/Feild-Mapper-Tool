"""
Script to update functional document to focus only on executable usage
Removes all Python code running instructions
"""

from docx import Document
from datetime import datetime

def fix_exe_only():
    doc = Document('Field_Mapper_Tool_Functional_Document_Final_20251118_170530.docx')
    
    changes_made = 0
    paragraphs_to_remove = []
    
    print("\nUpdating document for executable-only usage...")
    print("="*60)
    
    # Track sections to remove or modify
    in_method2_section = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        
        # Remove "Method 2: Running from Source" section
        if 'method 2:' in text and 'running from source' in text:
            paragraphs_to_remove.append(i)
            in_method2_section = True
            print(f"  ❌ Removing: 'Method 2: Running from Source' section")
            continue
        
        # Remove content under Method 2 until next heading
        if in_method2_section:
            if para.style.name.startswith('Heading'):
                in_method2_section = False
            elif para.style.name == 'List Bullet':
                paragraphs_to_remove.append(i)
            elif 'python field_mapper.py' in text:
                paragraphs_to_remove.append(i)
        
        # Remove Python-related content
        python_phrases = [
            'python field_mapper.py',
            'running from source',
            'pip install',
            'requirements.txt',
            'python 3.7',
            'python code',
            'command prompt',
            'navigate to the project folder',
        ]
        
        for phrase in python_phrases:
            if phrase in text and i not in paragraphs_to_remove:
                paragraphs_to_remove.append(i)
                print(f"  ❌ Removing paragraph with: '{phrase}'")
                break
    
    # Remove marked paragraphs (in reverse to maintain indices)
    for i in sorted(paragraphs_to_remove, reverse=True):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            changes_made += 1
    
    # Update "Launching the Application" section
    for i, para in enumerate(doc.paragraphs):
        if para.text == '2.2 Launching the Application':
            # Find and update next paragraph
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if 'two ways' in next_para.text.lower():
                    # Update to single method
                    for run in next_para.runs:
                        run.text = ''
                    if next_para.runs:
                        next_para.runs[0].text = 'To launch the Field Mapper Tool:'
                    else:
                        next_para.add_run('To launch the Field Mapper Tool:')
                    changes_made += 1
                    print(f"  ✏️  Updated: 'Launching the Application' section")
                    break
    
    # Update "Method 1" to just "How to Launch"
    for para in doc.paragraphs:
        if 'Method 1:' in para.text and 'Using the Executable' in para.text:
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = 'How to Launch:'
                para.runs[0].bold = True
            else:
                run = para.add_run('How to Launch:')
                run.bold = True
            changes_made += 1
            print(f"  ✏️  Changed: 'Method 1' → 'How to Launch'")
    
    # Update system requirements to remove Python mention
    for para in doc.paragraphs:
        if 'Python 3.7 or higher (if running from source)' in para.text:
            # This is in a list, remove this specific item
            p_element = para._element
            p_element.getparent().remove(p_element)
            changes_made += 1
            print(f"  ❌ Removed: Python requirement from system requirements")
    
    # Update "Getting Help" section to remove README references with code
    for para in doc.paragraphs:
        original = para.text
        new_text = original
        
        # Remove code-related references
        replacements = {
            'Review README.md for additional documentation': 'Check the documentation folder',
            'Check QUICK_START.md for quick reference': 'Refer to this user manual',
        }
        
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
        
        if new_text != original:
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
            changes_made += 1
            print(f"  ✏️  Updated help reference")
    
    # Save updated version
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'Field_Mapper_Tool_User_Manual_{timestamp}.docx'
    doc.save(filename)
    
    print("="*60)
    print(f"\n✅ Document updated successfully!")
    print(f"   Changes made: {changes_made}")
    print(f"   New filename: {filename}")
    print(f"\n   Document now focuses on:")
    print(f"   ✅ Executable (.exe) usage only")
    print(f"   ❌ No Python code instructions")
    print(f"   ❌ No 'pip install' commands")
    print(f"   ❌ No 'Method 2: Running from Source'")
    print(f"   ✅ Clean, user-friendly instructions")
    print("="*60)
    
    return filename

if __name__ == "__main__":
    fix_exe_only()

