"""
Document Parser Module
Extracts database and table field information from Word documents
"""

import re
from typing import Dict, List, Optional, Tuple
import logging

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    def __init__(self):
        self.parsed_data = {}
    
    def parse_document(self, file_path: str) -> Dict[str, Dict[str, List[str]]]:
        """
        Parse Word document to extract database/table field information
        
        Returns:
            Dictionary structure: {
                'database_name': {
                    'table_name': ['field1', 'field2', ...]
                }
            }
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            self.parsed_data = {}
            
            current_database = None
            current_table = None
            current_fields = []
            in_table_section = False
            table_index = 0  # Track table position in document
            
            # First, parse all paragraphs to understand structure
            paragraphs = list(doc.paragraphs)
            tables = list(doc.tables)
            
            # Parse document paragraph by paragraph
            for para_idx, paragraph in enumerate(paragraphs):
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # Look for database name patterns
                db_match = self._extract_database_name(text)
                if db_match:
                    # Save previous table if exists
                    if current_table and current_database:
                        if current_database not in self.parsed_data:
                            self.parsed_data[current_database] = {}
                        self.parsed_data[current_database][current_table] = current_fields.copy()
                    
                    current_database = db_match
                    current_table = None
                    current_fields = []
                    in_table_section = False
                    logger.info(f"Found database: {current_database}")
                    continue
                
                # Look for table/module name patterns
                table_match = self._extract_table_name(text)
                if table_match:
                    # Save previous table if exists
                    if current_table and current_database:
                        if current_database not in self.parsed_data:
                            self.parsed_data[current_database] = {}
                        self.parsed_data[current_database][current_table] = current_fields.copy()
                    
                    current_table = table_match
                    current_fields = []
                    in_table_section = True
                    logger.info(f"Found table/module: {current_table} in database: {current_database}")
                    continue
                
                # Check for "Data fields" indicator
                if re.search(r'Data fields[:\s–-]+', text, re.IGNORECASE):
                    in_table_section = True
                    # If no table name set, use a default based on database
                    if not current_table and current_database:
                        current_table = "Data Fields"
                        current_fields = []
                    continue
                
                # Extract field names if we're in a table section (from paragraphs)
                if in_table_section and current_database and current_table:
                    fields = self._extract_field_names(text)
                    if fields:
                        current_fields.extend(fields)
            
            # Now parse tables - they usually come after the paragraph structure
            for table_idx, table in enumerate(tables):
                # Try to find which database/table this belongs to
                # by looking at nearby paragraphs
                table_fields = self._extract_fields_from_table(table)
                
                if table_fields:
                    # Try to associate with current context
                    if current_database and current_table:
                        if current_database not in self.parsed_data:
                            self.parsed_data[current_database] = {}
                        if current_table not in self.parsed_data[current_database]:
                            self.parsed_data[current_database][current_table] = []
                        self.parsed_data[current_database][current_table].extend(table_fields)
                    else:
                        # Fallback: create generic entry
                        if 'Unknown' not in self.parsed_data:
                            self.parsed_data['Unknown'] = {}
                        table_name = f"Table_{table_idx + 1}"
                        if table_name not in self.parsed_data['Unknown']:
                            self.parsed_data['Unknown'][table_name] = []
                        self.parsed_data['Unknown'][table_name].extend(table_fields)
            
            # Save last table
            if current_table and current_database:
                if current_database not in self.parsed_data:
                    self.parsed_data[current_database] = {}
                if current_table not in self.parsed_data[current_database]:
                    self.parsed_data[current_database][current_table] = []
                self.parsed_data[current_database][current_table].extend(current_fields)
            
            logger.info(f"Parsed {len(self.parsed_data)} databases from document")
            return self.parsed_data
            
        except Exception as e:
            logger.error(f"Failed to parse document: {str(e)}")
            raise
    
    def _extract_fields_from_table(self, table) -> List[str]:
        """Extract field names from a Word table"""
        fields = []
        
        if len(table.rows) == 0:
            return fields
        
        # Get header row
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        
        # Look for field name column - prioritize "RDF Field Name" or "Field Name"
        field_column_idx = None
        
        # First, look for "RDF Field Name"
        for idx, header in enumerate(headers):
            if 'rdf field name' in header.lower():
                field_column_idx = idx
                break
        
        # If not found, look for "Field Name"
        if field_column_idx is None:
            for idx, header in enumerate(headers):
                if 'field name' in header.lower() and 'rdf' not in header.lower():
                    field_column_idx = idx
                    break
        
        # If still not found, look for any column with "field" or "name"
        if field_column_idx is None:
            for idx, header in enumerate(headers):
                if any(keyword in header.lower() for keyword in ['field', 'name', 'attribute']):
                    field_column_idx = idx
                    break
        
        # If no header found, assume second column (first is usually "No." or "Sr. No.")
        if field_column_idx is None and len(headers) > 1:
            field_column_idx = 1
        elif field_column_idx is None and len(headers) > 0:
            field_column_idx = 0
        
        # Extract fields from table
        if field_column_idx is not None:
            for row in table.rows[1:]:  # Skip header
                if len(row.cells) > field_column_idx:
                    field_text = row.cells[field_column_idx].text.strip()
                    field = self._clean_field_name(field_text)
                    if field and field not in fields:  # Avoid duplicates
                        fields.append(field)
        
        return fields
    
    def _extract_database_name(self, text: str) -> Optional[str]:
        """Extract database name from text"""
        # Patterns to match database names (including those with parentheses)
        patterns = [
            r'^([IVX]+\.\s*)?([A-Za-z0-9_\s\(\)]+?)\s+Database\s+details',  # "I. Database Name Database details"
            r'^([IVX]+\.\s*)?([A-Za-z0-9_\s\(\)]+?)\s+Database',  # "Database Name Database"
            r'Database[:\s]+([A-Za-z0-9_\s\(\)]+?)(?:\s+details|\s+are|\s|$|:)',  # "Database: Name" or "Database Name details"
            r'DB[:\s]+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',
            r'Database Name[:\s]+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',
            r'^([A-Za-z0-9_\s\(\)]+?)\s+Database',  # "Name Database"
            r'^([A-Za-z0-9_\s\(\)]+?)\s+DB',
            r'for\s+([A-Za-z0-9_\s\(\)]+?)\s+database',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                # Get the database name (could be in group 1 or 2 depending on pattern)
                db_name = match.group(2) if len(match.groups()) > 1 and match.group(2) else match.group(1)
                if db_name:
                    db_name = db_name.strip()
                    # Remove section numbers if present
                    db_name = re.sub(r'^[IVX]+\.\s*', '', db_name, flags=re.IGNORECASE)
                    # Clean up common suffixes and extra spaces
                    db_name = re.sub(r'\s+', ' ', db_name)
                    # Remove trailing "details" or "are"
                    db_name = re.sub(r'\s+(details|are|is)\s*$', '', db_name, flags=re.IGNORECASE)
                    if len(db_name) > 0 and len(db_name) < 150:  # Reasonable length
                        return db_name
        
        return None
    
    def _extract_table_name(self, text: str) -> Optional[str]:
        """Extract table name from text - can be module, sub-database, or table"""
        # Patterns to match table/module names
        patterns = [
            r'^([IVX]+\.\s*)?([A-Za-z0-9_\s\(\)]+?)\s+(Module|Database)',  # "I. Module Name Module" or "II. Database Name Database"
            r'Table[:\s]+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',
            r'^([A-Za-z0-9_\s\(\)]+?)\s+Table',
            r'Table Name[:\s]+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',
            r'Fields for[:\s]+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',
            r'Fields\s+in\s+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',
            r'Data fields[:\s–-]+([A-Za-z0-9_\s\(\)]+?)(?:\s|$|:)',  # "Data fields –" (with em dash)
            r'^([A-Za-z0-9_\s\(\)]+?)\s+fields',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                # Get the name (could be in group 1 or 2 depending on pattern)
                table_name = match.group(2) if len(match.groups()) > 1 and match.group(2) else match.group(1)
                if table_name:
                    table_name = table_name.strip()
                    # Remove section numbers if present
                    table_name = re.sub(r'^[IVX]+\.\s*', '', table_name, flags=re.IGNORECASE)
                    # Clean up common suffixes
                    table_name = re.sub(r'\s+(Module|Database|Table)\s*$', '', table_name, flags=re.IGNORECASE)
                    table_name = re.sub(r'\s+', ' ', table_name)
                    if len(table_name) > 0 and len(table_name) < 150:  # Reasonable length
                        return table_name
        
        return None
    
    def _extract_field_names(self, text: str) -> List[str]:
        """Extract field names from text"""
        fields = []
        
        # Pattern 1: Comma-separated list
        if ',' in text:
            parts = [p.strip() for p in text.split(',')]
            for part in parts:
                # Remove common prefixes/suffixes
                field = self._clean_field_name(part)
                if field and len(field) > 0:
                    fields.append(field)
        
        # Pattern 2: Bullet points or numbered list
        elif re.match(r'^[\d\-\•\*]\s+', text):
            field = self._clean_field_name(text)
            if field:
                fields.append(field)
        
        # Pattern 3: Single field name
        else:
            field = self._clean_field_name(text)
            if field and len(field) > 0 and not any(word in field.lower() for word in ['table', 'database', 'field', 'column']):
                fields.append(field)
        
        return fields
    
    def _clean_field_name(self, text: str) -> Optional[str]:
        """Clean and extract field name from text"""
        # Remove bullet points, numbers, and common prefixes
        text = re.sub(r'^[\d\-\•\*\s]+', '', text)
        text = re.sub(r'^(Field|Column|Name|Attribute)[:\s]+', '', text, flags=re.IGNORECASE)
        text = text.strip()
        
        # Remove trailing punctuation
        text = re.sub(r'[.,;:]+$', '', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Check if it looks like a field name (alphanumeric and underscores, possibly with spaces)
        # First try exact match
        if re.match(r'^[A-Za-z0-9_]+$', text):
            return text
        
        # Try with spaces (might be multi-word, take first word)
        if ' ' in text:
            first_word = text.split()[0]
            if re.match(r'^[A-Za-z0-9_]+$', first_word):
                return first_word
        
        return None
    
    def _parse_tables(self, doc: Document):
        """Parse tables in the Word document"""
        # We'll parse tables in context with the main parsing logic
        # This method is kept for backward compatibility but tables are now
        # parsed inline with the document structure
        pass
    
    def get_databases(self) -> List[str]:
        """Get list of all databases found in document"""
        return list(self.parsed_data.keys())
    
    def get_tables(self, database_name: str) -> List[str]:
        """Get list of tables for a specific database"""
        if database_name in self.parsed_data:
            return list(self.parsed_data[database_name].keys())
        return []
    
    def get_fields(self, database_name: str, table_name: str) -> List[str]:
        """Get field names for a specific database and table"""
        if database_name in self.parsed_data:
            if table_name in self.parsed_data[database_name]:
                return self.parsed_data[database_name][table_name]
        return []
    
    def get_all_fields_for_database(self, database_name: str) -> List[str]:
        """Get all fields across all tables for a database"""
        all_fields = []
        if database_name in self.parsed_data:
            for table_name, fields in self.parsed_data[database_name].items():
                all_fields.extend(fields)
        return list(set(all_fields))  # Remove duplicates

